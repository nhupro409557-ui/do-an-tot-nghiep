from __future__ import annotations

from datetime import datetime
from decimal import Decimal
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

SHOP_NAME = "ELECTROMART VIỆT NAM"
SHOP_DESCRIPTION = "Hệ thống bán lẻ điện thoại, laptop và phụ kiện công nghệ"
SHOP_ADDRESS_FALLBACK = ".............................................................."
SHOP_PHONE_FALLBACK = "..........................................................."
RECEIPT_REASON_LABELS = {
    "NK_MUA": "Nhập mua từ nhà cung cấp",
    "NK_TRA_NCC": "Nhà cung cấp trả lại hàng",
    "NK_KH_TRA": "Khách hàng trả hàng",
    "NK_BH": "Nhập bảo hành",
    "NK_DIEUCHINH": "Điều chỉnh tăng tồn kho",
    "NK_CHUYEN": "Nhập từ kho khác",
    "NK_SANXUAT": "Nhập thành phẩm",
    "NK_KHOI_TAO": "Nhập kho khởi tạo",
    "NK_KHAC": "Nhập khác",
}
PDF_FONT_DIR = Path(__file__).resolve().parents[2] / "assets" / "fonts"


def _register_pdf_fonts() -> tuple[str, str, str]:
    windows_fonts = Path("C:/Windows/Fonts")
    regular = windows_fonts / "arial.ttf"
    bold = windows_fonts / "arialbd.ttf"
    italic = windows_fonts / "ariali.ttf"
    if all(font_path.exists() for font_path in (regular, bold, italic)):
        font_paths = (regular, bold, italic, bold)
        family = "EMVArial"
    else:
        font_paths = (
            PDF_FONT_DIR / "Roboto-Regular.ttf",
            PDF_FONT_DIR / "Roboto-Bold.ttf",
            PDF_FONT_DIR / "Roboto-Italic.ttf",
            PDF_FONT_DIR / "Roboto-BoldItalic.ttf",
        )
        family = "EMVRoboto"
    try:
        normal_name = family
        bold_name = f"{family}-Bold"
        italic_name = f"{family}-Italic"
        bold_italic_name = f"{family}-BoldItalic"
        pdfmetrics.registerFont(TTFont(normal_name, str(font_paths[0])))
        pdfmetrics.registerFont(TTFont(bold_name, str(font_paths[1])))
        pdfmetrics.registerFont(TTFont(italic_name, str(font_paths[2])))
        pdfmetrics.registerFont(TTFont(bold_italic_name, str(font_paths[3])))
        pdfmetrics.registerFontFamily(
            family,
            normal=normal_name,
            bold=bold_name,
            italic=italic_name,
            boldItalic=bold_italic_name,
        )
        return normal_name, bold_name, italic_name
    except Exception:
        if family == "EMVArial":
            return _register_packaged_pdf_fonts()
        raise


def _register_packaged_pdf_fonts() -> tuple[str, str, str]:
    family = "EMVRoboto"
    faces = {
        family: PDF_FONT_DIR / "Roboto-Regular.ttf",
        f"{family}-Bold": PDF_FONT_DIR / "Roboto-Bold.ttf",
        f"{family}-Italic": PDF_FONT_DIR / "Roboto-Italic.ttf",
        f"{family}-BoldItalic": PDF_FONT_DIR / "Roboto-BoldItalic.ttf",
    }
    for name, path in faces.items():
        pdfmetrics.registerFont(TTFont(name, path))
    pdfmetrics.registerFontFamily(
        family,
        normal=family,
        bold=f"{family}-Bold",
        italic=f"{family}-Italic",
        boldItalic=f"{family}-BoldItalic",
    )
    return family, f"{family}-Bold", f"{family}-Italic"


def _num(value: Any) -> Decimal:
    if value is None or value == "":
        return Decimal("0")
    return Decimal(str(value))


def _int(value: Any) -> int:
    try:
        return int(value or 0)
    except (TypeError, ValueError):
        return 0


def _money(value: Any) -> str:
    return f"{int(round(float(_num(value)))):,}".replace(",", ".")


def _date_label(value: Any) -> str:
    if isinstance(value, datetime):
        date = value
    elif value:
        try:
            date = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        except ValueError:
            date = datetime.now()
    else:
        date = datetime.now()
    return f"Ngày {date.day:02d} tháng {date.month:02d} năm {date.year}"


def _receipt_reason(code: str | None) -> str:
    normalized = str(code or "NK_MUA").upper()
    label = RECEIPT_REASON_LABELS.get(normalized, RECEIPT_REASON_LABELS["NK_MUA"])
    return f"{normalized} - {label}"


def _variant_description(line: dict) -> str:
    parts = [
        str(line.get("variantColor") or "").strip(),
        str(line.get("variantConfiguration") or "").strip(),
    ]
    return " - ".join(part for part in parts if part)


def _actor(value: Any, label: Any = None) -> str:
    if label:
        return str(label)
    if value:
        text = str(value)
        return text[:8]
    return "-"


def _store_text(store_info: dict | None, key: str, fallback: str) -> str:
    if not store_info:
        return fallback
    value = str(store_info.get(key) or "").strip()
    return value or fallback


def _shop_header(store_info: dict | None = None) -> tuple[str, str, str, str]:
    return (
        _store_text(store_info, "name", SHOP_NAME),
        _store_text(store_info, "description", SHOP_DESCRIPTION),
        _store_text(store_info, "address", SHOP_ADDRESS_FALLBACK),
        _store_text(store_info, "hotline", SHOP_PHONE_FALLBACK),
    )


def _safe_filename(value: str, extension: str) -> str:
    cleaned = "".join(ch if ch.isalnum() or ch in "-_" else "-" for ch in value).strip("-")
    return f"{cleaned or 'phieu-nhap-kho'}.{extension}"


def amount_in_vietnamese(value: Any) -> str:
    units = ["", "nghìn", "triệu", "tỷ"]
    digits = ["không", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]

    def read_triple(num: int, full: bool = False) -> str:
        hundred = num // 100
        ten = (num % 100) // 10
        one = num % 10
        parts: list[str] = []
        if hundred > 0 or full:
            parts.append(f"{digits[hundred]} trăm")
        if ten > 1:
            parts.append(f"{digits[ten]} mươi")
            if one == 1:
                parts.append("mốt")
            elif one == 5:
                parts.append("lăm")
            elif one > 0:
                parts.append(digits[one])
        elif ten == 1:
            parts.append("mười")
            if one == 5:
                parts.append("lăm")
            elif one > 0:
                parts.append(digits[one])
        elif one > 0:
            if hundred > 0 or full:
                parts.append("lẻ")
            parts.append(digits[one])
        return " ".join(parts)

    rounded = int(round(float(_num(value))))
    if rounded <= 0:
        return "Không đồng."
    groups: list[int] = []
    current = rounded
    while current > 0:
        groups.append(current % 1000)
        current //= 1000
    words: list[str] = []
    for index, group in enumerate(groups):
        if group == 0:
            continue
        full = index < len(groups) - 1 and group < 100
        words.append(f"{read_triple(group, full)} {units[index]}".strip())
    result = " ".join(reversed(words))
    return f"{result[:1].upper()}{result[1:]} đồng."


def receipt_line_summaries(receipt: dict) -> list[dict]:
    summaries_by_key: dict[tuple[Any, ...], dict] = {}
    for line in receipt.get("lines") or []:
        planned = _int(line.get("plannedQuantity") or line.get("quantity"))
        tracks_imei = bool(line.get("tracksImei"))
        tracks_serial = bool(line.get("tracksSerialNumber"))
        imeis = line.get("imeis") if isinstance(line.get("imeis"), list) else []
        serial_numbers = line.get("serialNumbers") if isinstance(line.get("serialNumbers"), list) else []
        identifier_counts: list[int] = []
        if tracks_imei:
            identifier_counts.append(len(imeis))
        if tracks_serial:
            identifier_counts.append(len(serial_numbers))
        received = min(identifier_counts) if identifier_counts else _int(line.get("receivedQuantity") or planned)
        if received <= 0 and not identifier_counts:
            received = planned
        unit_cost = _num(line.get("unitCost"))
        identifier_status = "Không quản lý mã định danh"
        if tracks_imei or tracks_serial:
            parts: list[str] = []
            if tracks_imei:
                parts.append("Đủ IMEI" if len(imeis) >= planned else f"Thiếu {planned - len(imeis)} IMEI")
            if tracks_serial:
                parts.append(
                    "Đủ Serial"
                    if len(serial_numbers) >= planned
                    else f"Thiếu {planned - len(serial_numbers)} Serial"
                )
            identifier_status = " / ".join(parts)
        export_key = (
            line.get("productName"),
            line.get("variantSku"),
            line.get("variantColor"),
            line.get("variantConfiguration"),
            line.get("unitName") or "Cái",
            str(unit_cost),
        )
        existing = summaries_by_key.get(export_key)
        if existing:
            existing["planned"] += planned
            existing["received"] += received
            existing["amount"] += unit_cost * Decimal(received)
            existing["tracks_identifier"] = bool(existing["tracks_identifier"] or tracks_imei or tracks_serial)
            if identifier_status != "Không quản lý mã định danh":
                existing["identifier_status"] = identifier_status
            continue
        summaries_by_key[export_key] = {
            "line": line,
            "planned": planned,
            "received": received,
            "unit_cost": unit_cost,
            "amount": unit_cost * Decimal(received),
            "tracks_identifier": tracks_imei or tracks_serial,
            "identifier_status": identifier_status,
        }
    return list(summaries_by_key.values())


def _receipt_totals(summaries: list[dict]) -> tuple[int, int, Decimal]:
    planned = sum(item["planned"] for item in summaries)
    received = sum(item["received"] for item in summaries)
    total = sum((item["amount"] for item in summaries), Decimal("0"))
    return planned, received, total


def render_inventory_receipt_pdf(receipt: dict, store_info: dict | None = None) -> tuple[bytes, str]:
    normal_font, bold_font, italic_font = _register_pdf_fonts()
    summaries = receipt_line_summaries(receipt)
    total_planned, total_received, total_amount = _receipt_totals(summaries)
    shop_name, shop_description, shop_address, shop_phone = _shop_header(store_info)
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("ReceiptBody", parent=styles["Normal"], fontName=normal_font, fontSize=9, leading=12)
    small = ParagraphStyle("ReceiptSmall", parent=body, fontSize=8, leading=10)
    center = ParagraphStyle("ReceiptCenter", parent=body, alignment=TA_CENTER)
    right = ParagraphStyle("ReceiptRight", parent=body, alignment=TA_RIGHT)
    title = ParagraphStyle("ReceiptTitle", parent=center, fontName=bold_font, fontSize=18, leading=22)
    bold = ParagraphStyle("ReceiptBold", parent=body, fontName=bold_font)
    italic = ParagraphStyle("ReceiptItalic", parent=body, fontName=italic_font)

    receipt_date = _date_label(receipt.get("postedAt") or receipt.get("createdAt"))
    status = str(receipt.get("status") or "COMPLETED")
    is_official = status == "COMPLETED"

    story: list[Any] = [
        Table(
            [
                [
                    [
                        Paragraph(xml_escape(shop_name), bold),
                        Paragraph(xml_escape(shop_description), body),
                        Paragraph(f"Địa chỉ: {xml_escape(shop_address)}", body),
                        Paragraph(f"Điện thoại: {xml_escape(shop_phone)}", body),
                    ]
                ]
            ],
            colWidths=[173 * mm],
            style=[("VALIGN", (0, 0), (-1, -1), "TOP"), ("BOX", (0, 0), (-1, -1), 0, colors.white)],
        ),
        Spacer(1, 8),
        Paragraph("PHIẾU NHẬP KHO", title),
        Paragraph(receipt_date, ParagraphStyle("DateCenter", parent=italic, alignment=TA_CENTER, fontName=bold_font)),
        Paragraph(f"Số: {receipt.get('referenceCode') or '-'}", center),
    ]
    if not is_official:
        story.append(Paragraph(f"Phiếu nhập tạm - {status}", center))
    story.extend(
        [
            Spacer(1, 8),
            Paragraph(f"- Người giao hàng / Nhà cung cấp: <b>{receipt.get('supplierName') or '-'}</b>", body),
            Paragraph(f"- Theo chứng từ / Lý do nhập: {_receipt_reason(receipt.get('receiptReasonCode'))}", body),
            Paragraph(f"- Nhập tại kho: <b>{receipt.get('locationName') or 'Kho chính'}</b> &nbsp;&nbsp;&nbsp;&nbsp; Địa điểm: ........................................................", body),
            Paragraph(f"- Ghi chú: {receipt.get('note') or '-'}", body),
            Spacer(1, 5),
        ]
    )

    table_data: list[list[Any]] = [
        [
            Paragraph("STT", center),
            Paragraph("Tên, nhãn hiệu, quy cách phẩm chất sản phẩm, hàng hóa", center),
            Paragraph("Mã số", center),
            Paragraph("Đơn vị tính", center),
            Paragraph("Theo chứng từ", center),
            Paragraph("Thực nhập", center),
            Paragraph("Đơn giá", center),
            Paragraph("Thành tiền", center),
        ]
    ]
    for index, summary in enumerate(summaries, start=1):
        line = summary["line"]
        product_text = xml_escape(str(line.get("productName") or "-"))
        variant_description = _variant_description(line)
        if variant_description:
            product_text += f"<br/><font size='8'>{xml_escape(variant_description)}</font>"
        table_data.append(
            [
                Paragraph(str(index), center),
                Paragraph(product_text, body),
                Paragraph(str(line.get("variantSku") or line.get("productSku") or line.get("sku") or "-"), small),
                Paragraph(str(line.get("unitName") or "Cái"), center),
                Paragraph(_money(summary["planned"]), right),
                Paragraph(_money(summary["received"]), right),
                Paragraph(_money(summary["unit_cost"]), right),
                Paragraph(_money(summary["amount"]), right),
            ]
        )
    table_data.append(
        [
            "",
            Paragraph("Cộng", ParagraphStyle("TotalCenter", parent=center, fontName=bold_font)),
            "",
            "",
            Paragraph(_money(total_planned), ParagraphStyle("TotalRight1", parent=right, fontName=bold_font)),
            Paragraph(_money(total_received), ParagraphStyle("TotalRight2", parent=right, fontName=bold_font)),
            "",
            Paragraph(_money(total_amount), ParagraphStyle("TotalRight3", parent=right, fontName=bold_font)),
        ]
    )
    table = Table(
        table_data,
        repeatRows=1,
        colWidths=[12 * mm, 61 * mm, 24 * mm, 17 * mm, 19 * mm, 19 * mm, 24 * mm, 26 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.6, colors.black),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                ("FONTNAME", (0, 0), (-1, 0), bold_font),
                ("FONTNAME", (0, -1), (-1, -1), bold_font),
                ("ALIGN", (0, 0), (0, -1), "CENTER"),
            ]
        )
    )
    story.extend(
        [
            table,
            Spacer(1, 8),
            Paragraph(f"- Tổng số tiền (Viết bằng chữ): <b><i>{amount_in_vietnamese(total_amount)}</i></b>", body),
            Paragraph("- Số chứng từ gốc kèm theo: ..............................................................", body),
            Spacer(1, 8),
            Paragraph(receipt_date, ParagraphStyle("SignatureDate", parent=italic, alignment=TA_RIGHT)),
            Spacer(1, 6),
        ]
    )
    signatures = Table(
        [
            [
                Paragraph("<b>Người lập phiếu</b><br/><i>(Ký, họ tên)</i>", center),
                Paragraph("<b>Người giao hàng</b><br/><i>(Ký, họ tên)</i>", center),
                Paragraph("<b>Thủ kho</b><br/><i>(Ký, họ tên)</i>", center),
                Paragraph("<b>Kế toán trưởng</b><br/><i>(Hoặc bộ phận có nhu cầu nhập)</i><br/><i>(Ký, họ tên)</i>", center),
            ],
            [
                Paragraph(_actor(receipt.get("createdBy"), receipt.get("createdByName")), center),
                Paragraph(str(receipt.get("supplierName") or ""), center),
                Paragraph(_actor(receipt.get("postedBy"), receipt.get("postedByName")), center),
                Paragraph(_actor(receipt.get("approvedBy"), receipt.get("approvedByName")), center),
            ],
        ],
        colWidths=[43 * mm, 43 * mm, 43 * mm, 43 * mm],
        rowHeights=[26 * mm, 10 * mm],
    )
    signatures.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(signatures)
    doc.build(story)
    return buffer.getvalue(), _safe_filename(str(receipt.get("referenceCode") or "phieu-nhap-kho"), "pdf")


def _docx_text(cell, text: str = "", bold: bool = False, italic: bool = False, align=None) -> None:
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def render_inventory_receipt_docx(receipt: dict, store_info: dict | None = None) -> tuple[bytes, str]:
    summaries = receipt_line_summaries(receipt)
    total_planned, total_received, total_amount = _receipt_totals(summaries)
    receipt_date = _date_label(receipt.get("postedAt") or receipt.get("createdAt"))
    status = str(receipt.get("status") or "COMPLETED")
    shop_name, shop_description, shop_address, shop_phone = _shop_header(store_info)
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    header = document.add_table(rows=1, cols=1)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    _docx_text(header.cell(0, 0), f"{shop_name}\n{shop_description}\nĐịa chỉ: {shop_address}\nĐiện thoại: {shop_phone}", bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PHIẾU NHẬP KHO")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    date_p = document.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_p.add_run(receipt_date)
    date_run.bold = True
    date_run.italic = True
    document.add_paragraph(f"Số: {receipt.get('referenceCode') or '-'}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    if status != "COMPLETED":
        temp_p = document.add_paragraph(f"Phiếu nhập tạm - {status}")
        temp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"- Người giao hàng / Nhà cung cấp: {receipt.get('supplierName') or '-'}")
    document.add_paragraph(f"- Theo chứng từ / Lý do nhập: {_receipt_reason(receipt.get('receiptReasonCode'))}")
    document.add_paragraph(f"- Nhập tại kho: {receipt.get('locationName') or 'Kho chính'}    Địa điểm: ........................................................")
    document.add_paragraph(f"- Ghi chú: {receipt.get('note') or '-'}")

    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["STT", "Tên, nhãn hiệu, quy cách phẩm chất sản phẩm, hàng hóa", "Mã số", "Đơn vị tính", "Theo chứng từ", "Thực nhập", "Đơn giá", "Thành tiền"]
    for index, header_text in enumerate(headers):
        _docx_text(table.rows[0].cells[index], header_text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for index, summary in enumerate(summaries, start=1):
        line = summary["line"]
        row = table.add_row().cells
        product_text = str(line.get("productName") or "-")
        variant_description = _variant_description(line)
        if variant_description:
            product_text += f"\n{variant_description}"
        values = [
            str(index),
            product_text,
            str(line.get("variantSku") or line.get("productSku") or line.get("sku") or "-"),
            str(line.get("unitName") or "Cái"),
            _money(summary["planned"]),
            _money(summary["received"]),
            _money(summary["unit_cost"]),
            _money(summary["amount"]),
        ]
        for cell_index, value in enumerate(values):
            align = WD_ALIGN_PARAGRAPH.RIGHT if cell_index >= 4 else WD_ALIGN_PARAGRAPH.CENTER if cell_index in {0, 3} else WD_ALIGN_PARAGRAPH.LEFT
            _docx_text(row[cell_index], value, align=align)
    total_row = table.add_row().cells
    totals = ["", "Cộng", "", "", _money(total_planned), _money(total_received), "", _money(total_amount)]
    for cell_index, value in enumerate(totals):
        align = WD_ALIGN_PARAGRAPH.RIGHT if cell_index in {4, 5, 7} else WD_ALIGN_PARAGRAPH.CENTER
        _docx_text(total_row[cell_index], value, bold=True, align=align)

    document.add_paragraph(f"- Tổng số tiền (Viết bằng chữ): {amount_in_vietnamese(total_amount)}")
    document.add_paragraph("- Số chứng từ gốc kèm theo: ..............................................................")
    date_sign = document.add_paragraph(receipt_date)
    date_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature = document.add_table(rows=2, cols=4)
    signature.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        "Người lập phiếu\n(Ký, họ tên)",
        "Người giao hàng\n(Ký, họ tên)",
        "Thủ kho\n(Ký, họ tên)",
        "Kế toán trưởng\n(Hoặc bộ phận có nhu cầu nhập)\n(Ký, họ tên)",
    ]
    names = [
        _actor(receipt.get("createdBy"), receipt.get("createdByName")),
        str(receipt.get("supplierName") or ""),
        _actor(receipt.get("postedBy"), receipt.get("postedByName")),
        _actor(receipt.get("approvedBy"), receipt.get("approvedByName")),
    ]
    for index, label in enumerate(labels):
        _docx_text(signature.rows[0].cells[index], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _docx_text(signature.rows[1].cells[index], names[index], align=WD_ALIGN_PARAGRAPH.CENTER)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue(), _safe_filename(str(receipt.get("referenceCode") or "phieu-nhap-kho"), "docx")


def render_order_invoice_pdf(order: Any, items: list[Any]) -> tuple[bytes, str]:
    normal_font, bold_font, italic_font = _register_pdf_fonts()
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("InvoiceBody", parent=styles["Normal"], fontName=normal_font, fontSize=9, leading=12)
    center = ParagraphStyle("InvoiceCenter", parent=body, alignment=TA_CENTER)
    right = ParagraphStyle("InvoiceRight", parent=body, alignment=TA_RIGHT)
    title = ParagraphStyle("InvoiceTitle", parent=center, fontName=bold_font, fontSize=18, leading=22)
    bold = ParagraphStyle("InvoiceBold", parent=body, fontName=bold_font)
    italic = ParagraphStyle("InvoiceItalic", parent=body, fontName=italic_font)

    def get_val(obj: Any, key: str, default: Any = None) -> Any:
        if isinstance(obj, dict):
            return obj.get(key, default)
        return getattr(obj, key, default)

    order_code = get_val(order, "order_code") or get_val(order, "orderCode", "-")
    created_at = get_val(order, "created_at") or get_val(order, "createdAt")
    invoice_date = _date_label(get_val(order, "completed_at") or get_val(order, "completedAt") or created_at)

    recipient_name = get_val(order, "recipient_name") or get_val(order, "recipientName", "-")
    recipient_phone = get_val(order, "recipient_phone") or get_val(order, "recipientPhone", "-")
    shipping_address = get_val(order, "shipping_address") or get_val(order, "shippingAddress", "-")
    payment_method = get_val(order, "payment_method") or get_val(order, "paymentMethod", "-")
    payment_status = get_val(order, "payment_status") or get_val(order, "paymentStatus", "-")
    subtotal = _num(get_val(order, "subtotal_amount") or get_val(order, "subtotalAmount", 0))
    discount = _num(get_val(order, "discount_amount") or get_val(order, "discountAmount", 0))
    shipping_fee = _num(get_val(order, "shipping_fee") or get_val(order, "shippingFee", 0))
    total_amount = _num(get_val(order, "total_amount") or get_val(order, "totalAmount", 0))

    story: list[Any] = [
        Table(
            [
                [
                    [
                        Paragraph(SHOP_NAME, bold),
                        Paragraph(SHOP_DESCRIPTION, body),
                        Paragraph("Địa chỉ: 123 Đường Công Nghệ, Quận 1, TP. Hồ Chí Minh", body),
                        Paragraph("Điện thoại: 1900 1234 - Email: support@electromart.vn", body),
                    ]
                ]
            ],
            colWidths=[173 * mm],
            style=[("VALIGN", (0, 0), (-1, -1), "TOP"), ("BOX", (0, 0), (-1, -1), 0, colors.white)],
        ),
        Spacer(1, 8),
        Paragraph("HÓA ĐƠN BÁN HÀNG", title),
        Paragraph(invoice_date, ParagraphStyle("DateCenter", parent=italic, alignment=TA_CENTER, fontName=bold_font)),
        Paragraph(f"Mã đơn hàng: {order_code}", center),
        Spacer(1, 8),
        Paragraph(f"- Khách hàng: <b>{recipient_name}</b> &nbsp;&nbsp;&nbsp;&nbsp; Điện thoại: <b>{recipient_phone}</b>", body),
        Paragraph(f"- Địa chỉ nhận hàng: {shipping_address}", body),
        Paragraph(f"- Phương thức thanh toán: <b>{payment_method}</b> &nbsp;&nbsp;&nbsp;&nbsp; Trạng thái: <b>{payment_status}</b>", body),
        Spacer(1, 5),
    ]

    table_data: list[list[Any]] = [
        [
            Paragraph("STT", center),
            Paragraph("Tên sản phẩm", center),
            Paragraph("Đơn vị tính", center),
            Paragraph("Số lượng", center),
            Paragraph("Đơn giá (VND)", center),
            Paragraph("Thành tiền (VND)", center),
        ]
    ]

    for index, item in enumerate(items, start=1):
        item_name = get_val(item, "product_name") or get_val(item, "productName", "-")
        qty = _int(get_val(item, "quantity", 1))
        price = _num(get_val(item, "unit_price") or get_val(item, "price", 0))
        item_total = _num(get_val(item, "total_price") or get_val(item, "totalPrice", price * qty))

        table_data.append(
            [
                Paragraph(str(index), center),
                Paragraph(item_name, body),
                Paragraph("Cái", center),
                Paragraph(str(qty), center),
                Paragraph(_money(price), right),
                Paragraph(_money(item_total), right),
            ]
        )

    table_data.extend([
        [
            "", Paragraph("Cộng tiền hàng", bold), "", "", "",
            Paragraph(_money(subtotal), ParagraphStyle("BoldRight", parent=right, fontName=bold_font))
        ],
        [
            "", Paragraph("Giảm giá (Voucher / Điểm)", bold), "", "", "",
            Paragraph(f"-{_money(discount)}", ParagraphStyle("BoldRight", parent=right, fontName=bold_font))
        ],
        [
            "", Paragraph("Phí vận chuyển", bold), "", "", "",
            Paragraph(_money(shipping_fee), ParagraphStyle("BoldRight", parent=right, fontName=bold_font))
        ],
        [
            "", Paragraph("Tổng thanh toán", bold), "", "", "",
            Paragraph(_money(total_amount), ParagraphStyle("TotalRight", parent=right, fontName=bold_font, textColor=colors.HexColor("#d70018")))
        ]
    ])

    col_widths = [10 * mm, 78 * mm, 20 * mm, 15 * mm, 25 * mm, 25 * mm]
    table_style = TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f3f4f6")),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
            ("GRID", (0, 0), (-1, -5), 0.5, colors.HexColor("#d1d5db")),
            ("BOX", (0, -4), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
            ("LINEBELOW", (0, -4), (-1, -1), 0.5, colors.HexColor("#e5e7eb")),
            ("SPAN", (1, -4), (4, -4)),
            ("SPAN", (1, -3), (4, -3)),
            ("SPAN", (1, -2), (4, -2)),
            ("SPAN", (1, -1), (4, -1)),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]
    )

    story.append(Table(table_data, colWidths=col_widths, style=table_style))
    story.extend([
        Spacer(1, 8),
        Paragraph(f"- Tổng tiền thanh toán (viết bằng chữ): <b><i>{amount_in_vietnamese(total_amount)}</i></b>", body),
        Spacer(1, 15),
    ])

    signature_data = [
        [
            Paragraph("<b>Khách hàng</b><br/>(Ký, ghi rõ họ tên)", center),
            Paragraph("<b>Người bán hàng</b><br/>(Ký, ghi rõ họ tên)", center),
        ],
        ["", ""],
    ]
    story.append(
        Table(
            signature_data,
            colWidths=[86 * mm, 87 * mm],
            style=[
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 40),
            ],
        )
    )

    doc.build(story)
    return buffer.getvalue(), f"hoa-don-{order_code}.pdf"
